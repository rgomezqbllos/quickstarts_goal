#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
goalbus_pdf_pt.py
─────────────────
Convierte archivos Markdown de goalbus (ya en PORTUGUÉS) a PDFs imprimibles.

REQUISITO (instalar una sola vez):
    pip install reportlab

USO — carpeta completa:
    python3 goalbus_pdf_pt.py --md-dir "/ruta/a/markdowns_pt/"

USO — archivo individual:
    python3 goalbus_pdf_pt.py --md "/ruta/a/P10_foo.md"

OPCIONES:
    --md-dir   Carpeta con archivos Pxx_*.md en portugués
    --md       Archivo markdown individual
    --logo     Ruta a goal-logo-white.png (opcional; se busca automáticamente)
    --out      Carpeta de salida (por defecto: misma que --md-dir o --md)
    --from N   Procesar desde el número N (por defecto: 1)
    --to   N   Procesar hasta el número N (por defecto: 99)

SALIDA:
    P01_QuickStart_Imprimivel.pdf, P02_..., etc.
"""

import sys, os, re, io, math, argparse

try:
    from PIL import Image as PILImage
except Exception:
    PILImage = None

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.lib.colors import HexColor
    from reportlab.lib.utils import ImageReader
    from reportlab.platypus import SimpleDocTemplate, Spacer, KeepTogether, Table, TableStyle
    from reportlab.platypus.flowables import Flowable
    from reportlab.platypus.paragraph import Paragraph
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
except ImportError:
    print("ERROR: Falta 'reportlab'. Instala con:  pip install reportlab")
    sys.exit(1)

# ── Colores ─────────────────────────────────────────────────────────────────
BG_PAGE    = HexColor("#0D1520"); BG_CARD    = HexColor("#141F30")
BG_CARD_ALT= HexColor("#1A2640"); BG_SURFACE = HexColor("#0A1018")
TEAL       = HexColor("#6EF8EC"); TEAL_DIM   = HexColor("#2BC4BA")
BLUE_LIGHT = HexColor("#60A5FA"); TEXT_WHITE = HexColor("#F0F4F8")
TEXT_GRAY  = HexColor("#94A3B8"); TEXT_DIM   = HexColor("#64748B")
WARN       = HexColor("#FBBF24"); SUCCESS    = HexColor("#34D399")
W, H = letter; PAD = 0.45 * inch
TOP_MARGIN = 0.75 * inch
BOTTOM_MARGIN = 0.6 * inch
CONTENT_MAX_IMAGE_FRAC = 0.70
_LOGO_RATIO = 1868 / 1031
LOGO_PATH = ""  # se asigna en main()
ALLOWED_IMAGE_EXTS = (".png", ".jpg", ".jpeg")
REF_LINE_RE = re.compile(r'^\s*ref:\s*(.*?)\s*$', re.IGNORECASE)
OUTPUT_FORMATS = ("pdf", "docx")
IMAGE_SIZE_OPTIONS = ("auto", "full", "compact")
DEFAULT_IMAGE_SIZE = "auto"
DEFAULT_COMPACT_WIDTH_RATIO = 0.68
DEFAULT_COMPACT_HEIGHT_FRAC = 0.36
DEFAULT_COMPACT_HEIGHT_FRAC_VERTICAL = 0.30
DEFAULT_COMPACT_HEIGHT_FRAC_PANORAMIC = 0.42
MAX_IMAGE_SPLIT_PARTS = 6
IMAGE_FRAME_COLOR = HexColor("#2A3340")
IMAGE_FRAME_LINE_WIDTH = 0.7
IMAGE_SHADOW_OFFSET = 0.028 * inch
IMAGE_SHADOW_ALPHA = 0.10

# Valores hex para salida DOCX (Modo Oscuro Restaurado)
DOCX_BG_PAGE = "0D1520"
DOCX_BG_CARD = "141F30"
DOCX_BG_CARD_ALT = "1A2640"
DOCX_BG_SURFACE = "0A1018"
DOCX_TEAL = "6EF8EC"
DOCX_TEAL_DIM = "2BC4BA"
DOCX_BLUE_LIGHT = "60A5FA"
DOCX_TEXT_WHITE = "F0F4F8"
DOCX_TEXT_GRAY = "94A3B8"
DOCX_TEXT_DIM = "64748B"
DOCX_WARN = "FBBF24"
DOCX_SUCCESS = "34D399"
DOCX_IMAGE_FRAME = "2A3340"
DOCX_IMAGE_FRAME_WIDTH = "7620"
DOCX_IMAGE_SHADOW_ALPHA = "9000"
DOCX_IMAGE_SHADOW_BLUR = "63500"
DOCX_IMAGE_SHADOW_DIST = "38100"

# ── Fuentes (detección automática: macOS / Linux / Windows) ─────────────────
def _setup_fonts():
    def try_reg(name, paths):
        for p in paths:
            if os.path.exists(p):
                try:
                    pdfmetrics.registerFont(TTFont(name, p))
                    return True
                except Exception:
                    pass
        return False

    regular = [
        "/System/Library/Fonts/Supplemental/Arial.ttf",  # macOS Catalina+
        "/Library/Fonts/Arial.ttf",
        os.path.expanduser("~/Library/Fonts/Arial.ttf"),
        "/System/Library/Fonts/Supplemental/Verdana.ttf",
        "/Library/Fonts/Verdana.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",      # Linux
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        "C:/Windows/Fonts/arial.ttf",                           # Windows
        "C:/Windows/Fonts/segoeui.ttf",
    ]
    bold = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf",
        "/Library/Fonts/Arial Bold.ttf",
        os.path.expanduser("~/Library/Fonts/Arial Bold.ttf"),
        "/System/Library/Fonts/Supplemental/Verdana Bold.ttf",
        "/Library/Fonts/Verdana Bold.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
        "C:/Windows/Fonts/arialbd.ttf",
        "C:/Windows/Fonts/segoeuib.ttf",
    ]

    ok_r = try_reg("Inter", regular)
    ok_b = try_reg("Inter-Bold", bold)

    # Último recurso: escanear carpetas del sistema
    if not ok_r or not ok_b:
        for d in ["/System/Library/Fonts/Supplemental/", "/Library/Fonts/",
                  os.path.expanduser("~/Library/Fonts/"), "/usr/share/fonts/"]:
            if not os.path.isdir(d): continue
            for f in sorted(os.listdir(d)):
                if not f.lower().endswith(".ttf"): continue
                p = os.path.join(d, f)
                if not ok_r:
                    try: pdfmetrics.registerFont(TTFont("Inter", p)); ok_r = True
                    except Exception: pass
                if not ok_b and ok_r:
                    try: pdfmetrics.registerFont(TTFont("Inter-Bold", p)); ok_b = True
                    except Exception: pass
                if ok_r and ok_b: break
            if ok_r and ok_b: break

_setup_fonts()

# ── Utilidades de texto ──────────────────────────────────────────────────────
def parse_bold(text):
    parts, i, r = text.split("**"), 0, []
    for p in parts: r.append((p, i % 2 == 1)); i += 1
    return r

def measure_wrapped_bold(text, size, max_w):
    from reportlab.pdfbase.pdfmetrics import stringWidth
    words = [(w, b) for seg, b in parse_bold(text) for w in seg.split()]
    lh = size * 1.45; lines = 1; cur_w = 0
    for w, b in words:
        f = "Inter-Bold" if b else "Inter"
        ww = stringWidth(w, f, size); sp = stringWidth(" ", "Inter", size)
        if cur_w == 0: cur_w = ww
        elif cur_w + sp + ww <= max_w: cur_w += sp + ww
        else: lines += 1; cur_w = ww
    return lines * lh

def draw_wrapped_bold(c, text, x, start_y, size, max_w, cn, cb):
    words = [(w, b) for seg, b in parse_bold(text) for w in seg.split()]
    lh = size * 1.45; y = start_y; cur = []
    def lw(ws): return sum(c.stringWidth(w, "Inter-Bold" if b else "Inter", size) + (c.stringWidth(" ", "Inter", size) if i < len(ws)-1 else 0) for i, (w, b) in enumerate(ws))
    def flush(ws, yy):
        cx = x
        for w, b in ws:
            f = "Inter-Bold" if b else "Inter"; c.setFillColor(cb if b else cn)
            c.setFont(f, size); c.drawString(cx, yy, w)
            cx += c.stringWidth(w, f, size) + c.stringWidth(" ", "Inter", size)
    for word, bold in words:
        test = cur + [(word, bold)]
        if lw(test) <= max_w: cur = test
        else:
            if cur: flush(cur, y); y -= lh
            cur = [(word, bold)]
    if cur: flush(cur, y); y -= lh
    return y

# ── Header / Footer ──────────────────────────────────────────────────────────
def make_bg(label, footer):
    top_pad = 8  # ~2.8 mm of aire sobre el logo
    def _bg(c, doc):
        c.saveState()
        c.setFillColor(BG_PAGE); c.rect(0, 0, W, H, fill=1, stroke=0)
        header_y = H-0.55*inch - top_pad
        c.setFillColor(BG_SURFACE); c.rect(0, header_y, W, 0.55*inch, fill=1, stroke=0)
        if LOGO_PATH and os.path.exists(LOGO_PATH):
            lw, lh2 = 1.1*inch, 1.1*inch/_LOGO_RATIO
            c.drawImage(LOGO_PATH, PAD, header_y+(0.55*inch-lh2)/2, width=lw, height=lh2, mask="auto")
        c.setFont("Inter", 7.5); c.setFillColor(TEXT_DIM)
        c.drawRightString(W-PAD, header_y+0.18*inch, label)
        c.setFillColor(BG_SURFACE); c.rect(0, 0, W, 0.38*inch, fill=1, stroke=0)
        c.setFont("Inter", 7); c.setFillColor(TEXT_DIM)
        c.drawString(PAD, 0.13*inch, footer)
        c.drawRightString(W-PAD, 0.13*inch, f"Página {doc.page}")
        c.restoreState()
    return _bg

# ── Flowables ────────────────────────────────────────────────────────────────
class HeroBlock(Flowable):
    def __init__(self, number, title, intro):
        super().__init__(); self.number = number; self.title = title; self.intro = intro; self.avail = W-2*PAD
    def _h(self):
        return 0.32*inch + measure_wrapped_bold(self.title, 18, self.avail-0.8*inch) + 0.12*inch + measure_wrapped_bold(self.intro, 9.5, self.avail-0.1*inch) + 0.28*inch
    def wrap(self, aw, ah): self.width = aw; self.height = self._h(); return aw, self.height
    def draw(self):
        c = self.canv; aw = self.avail
        c.setFillColor(BG_CARD); c.roundRect(0, 0, aw, self.height, 8, fill=1, stroke=0)
        c.setFillColor(TEAL_DIM); c.rect(0, 0, 4, self.height, fill=1, stroke=0)
        y = self.height-0.32*inch; c.setFont("Inter-Bold", 11); c.setFillColor(TEAL); c.drawString(0.18*inch, y, f"P{self.number}")
        y = draw_wrapped_bold(c, self.title, 0.18*inch, y-0.02*inch-18*1.2, 18, aw-0.8*inch, TEXT_WHITE, TEAL) - 0.06*inch
        draw_wrapped_bold(c, self.intro, 0.18*inch, y, 9.5, aw-0.3*inch, TEXT_GRAY, TEXT_WHITE)

class SectionHeader(Flowable):
    def __init__(self, text, number=None):
        super().__init__(); self.text = text; self.number = number; self.avail = W-2*PAD
    def wrap(self, aw, ah):
        self.width = aw
        max_w = self.avail-0.65*inch if self.number else self.avail-0.3*inch
        self.height = measure_wrapped_bold(self.text, 12.5, max_w) + 0.32*inch
        return aw, self.height
    def draw(self):
        c = self.canv; aw = self.avail
        c.setFillColor(BG_CARD_ALT); c.roundRect(0, 0, aw, self.height, 6, fill=1, stroke=0)
        c.setFillColor(TEAL); c.rect(0, 0, 3, self.height, fill=1, stroke=0)
        max_w = self.avail-0.65*inch if self.number else self.avail-0.3*inch
        text_h = measure_wrapped_bold(self.text, 12.5, max_w)
        text_y = (self.height + text_h)/2 - 12.5*0.25
        if self.number:
            c.setFont("Inter-Bold", 9); c.setFillColor(TEAL_DIM)
            num_y = (self.height - 9)/2
            c.drawString(0.18*inch, num_y, f"0{self.number}" if self.number < 10 else str(self.number))
            draw_wrapped_bold(c, self.text, 0.52*inch, text_y, 12.5, max_w, TEAL, TEAL)
        else:
            draw_wrapped_bold(c, self.text, 0.18*inch, text_y, 12.5, max_w, TEAL, TEAL)

class BodyText(Flowable):
    def __init__(self, text, size=9.5, color=None):
        super().__init__(); self.text = text; self.size = size; self.color = color or TEXT_GRAY; self.avail = W-2*PAD
    def wrap(self, aw, ah): self.width = aw; self.height = measure_wrapped_bold(self.text, self.size, self.avail); return aw, self.height
    def draw(self): draw_wrapped_bold(self.canv, self.text, 0, self.height-self.size*0.25, self.size, self.avail, self.color, TEXT_WHITE)

class PrereqBox(Flowable):
    def __init__(self, items):
        super().__init__(); self.items = items; self.avail = W-2*PAD
    def _h(self):
        h = 0.28*inch
        for it in self.items: h += measure_wrapped_bold("•  "+it, 9, self.avail-0.55*inch) + 4
        return h + 0.18*inch
    def wrap(self, aw, ah): self.width = aw; self.height = self._h(); return aw, self.height
    def draw(self):
        c = self.canv; aw = self.avail
        c.setFillColor(BG_SURFACE); c.roundRect(0, 0, aw, self.height, 6, fill=1, stroke=0)
        c.setStrokeColor(WARN); c.setLineWidth(1); c.roundRect(0, 0, aw, self.height, 6, fill=0, stroke=1)
        y = self.height-0.2*inch
        for it in self.items:
            ih = measure_wrapped_bold("•  "+it, 9, aw-0.55*inch); y -= ih
            draw_wrapped_bold(c, "•  "+it, 0.22*inch, y+ih-9*0.25, 9, aw-0.55*inch, TEXT_GRAY, TEXT_WHITE); y -= 4

class CaseRefBox(Flowable):
    def __init__(self, text):
        super().__init__(); self.text = text; self.avail = W-2*PAD
    def wrap(self, aw, ah):
        self.width = aw; self.height = measure_wrapped_bold(self.text, 9.5, self.avail-0.65*inch) + 0.52*inch; return aw, self.height
    def draw(self):
        c = self.canv; aw = self.avail
        c.setFillColor(BG_CARD); c.roundRect(0, 0, aw, self.height, 6, fill=1, stroke=0)
        c.setStrokeColor(BLUE_LIGHT); c.setLineWidth(1); c.roundRect(0, 0, aw, self.height, 6, fill=0, stroke=1)
        draw_wrapped_bold(c, self.text, 0.18*inch, self.height-0.2*inch, 9.5, aw-0.38*inch, TEXT_GRAY, TEXT_WHITE)

class NumberedStepsList(Flowable):
    def __init__(self, steps):
        super().__init__(); self.steps = steps; self.avail = W-2*PAD
    def _h(self):
        h = 0.22*inch; ind = PAD+26
        for s in self.steps:
            mw = self.avail-0.45*inch if s.get("snum", "") == "•" else self.avail-(ind-PAD)-0.22*inch
            h += measure_wrapped_bold(s["text"], 9, mw) + 5
            for sub in s.get("subs", []):
                sub_text = sub.get("text") if isinstance(sub, dict) else str(sub)
                h += measure_wrapped_bold(sub_text, 8.5, self.avail-0.85*inch) + 4
        return h + 0.12*inch
    def wrap(self, aw, ah): self.width = aw; self.height = self._h(); return aw, self.height
    def draw(self):
        c = self.canv; aw = self.avail
        c.setFillColor(BG_CARD); c.roundRect(0, 0, aw, self.height, 6, fill=1, stroke=0)
        ind = 0.36*inch; y = self.height-0.28*inch
        for s in self.steps:
            snum = s.get("snum", ""); is_b = (snum == "•")
            tx = ind+0.22*inch if not is_b else 0.35*inch
            mw = aw-tx-0.1*inch; th = measure_wrapped_bold(s["text"], 9, mw); y -= th
            fy = y+th-9*0.25
            if is_b:
                c.setFillColor(TEAL); c.setFont("Inter-Bold", 11); c.drawString(0.18*inch, fy, "•")
            else:
                bw = c.stringWidth(str(snum), "Inter-Bold", 8.5); cy = fy+9*0.35
                c.setFillColor(TEAL_DIM); c.circle(ind, cy, 8, fill=1, stroke=0)
                c.setFont("Inter-Bold", 8); c.setFillColor(BG_PAGE); c.drawString(ind-bw/2, cy-3.5, str(snum))
            draw_wrapped_bold(c, s["text"], tx, fy, 9, mw, TEXT_GRAY, TEXT_WHITE); y -= 5
            for sub in s.get("subs", []):
                if isinstance(sub, dict):
                    sub_text = sub.get("text", "")
                    sub_snum = sub.get("snum", "–")
                else:
                    sub_text = str(sub); sub_snum = "–"
                sh = measure_wrapped_bold(sub_text, 8.5, aw-0.75*inch); y -= sh; fy2 = y+sh-8.5*0.25
                c.setFont("Inter", 8.5); c.setFillColor(TEXT_DIM)
                if isinstance(sub_snum, int):
                    c.drawString(0.55*inch, fy2, f"{sub_snum}.")
                elif sub_snum == "•":
                    c.setFont("Inter-Bold", 9); c.setFillColor(TEAL_DIM)
                    c.drawString(0.55*inch, fy2, "•")
                    c.setFont("Inter", 8.5); c.setFillColor(TEXT_DIM)
                else:
                    c.drawString(0.55*inch, fy2, "–")
                draw_wrapped_bold(c, sub_text, 0.68*inch, fy2, 8.5, aw-0.78*inch, TEXT_DIM, TEXT_GRAY); y -= 4

class ExampleBox(Flowable):
    def __init__(self, title, items, numbered=True):
        super().__init__(); self.title = title; self.items = items; self.numbered = numbered; self.avail = W-2*PAD
    def _h(self):
        h = 0.32*inch
        for it in self.items: h += measure_wrapped_bold(it, 9, self.avail-0.55*inch) + 5
        return h + 0.12*inch
    def wrap(self, aw, ah): self.width = aw; self.height = self._h(); return aw, self.height
    def draw(self):
        c = self.canv; aw = self.avail
        c.setFillColor(BG_CARD_ALT); c.roundRect(0, 0, aw, self.height, 6, fill=1, stroke=0)
        c.setStrokeColor(SUCCESS); c.setLineWidth(0.8); c.roundRect(0, 0, aw, self.height, 6, fill=0, stroke=1)
        if self.title:
            c.setFont("Inter-Bold", 8); c.setFillColor(SUCCESS); c.drawString(0.18*inch, self.height-0.21*inch, self.title)
            y = self.height-0.36*inch
        else:
            y = self.height-0.2*inch
        for i, it in enumerate(self.items):
            ih = measure_wrapped_bold(it, 9, aw-0.55*inch); y -= ih; fy = y+ih-9*0.25
            if self.numbered:
                c.setFont("Inter-Bold", 8.5); c.setFillColor(SUCCESS); c.drawString(0.18*inch, fy, f"{i+1}.")
                draw_wrapped_bold(c, it, 0.38*inch, fy, 9, aw-0.52*inch, TEXT_GRAY, TEXT_WHITE)
            else:
                c.setFont("Inter-Bold", 11); c.setFillColor(SUCCESS); c.drawString(0.18*inch, fy, "•")
                draw_wrapped_bold(c, it, 0.35*inch, fy, 9, aw-0.5*inch, TEXT_GRAY, TEXT_WHITE)
            y -= 5

class QuestionList(Flowable):
    def __init__(self, title, items):
        super().__init__(); self.title = title; self.items = items; self.avail = W-2*PAD
    def _h(self):
        h = 0.32*inch
        for it in self.items: h += measure_wrapped_bold(it, 9, self.avail-0.55*inch) + 5
        return h + 0.12*inch
    def wrap(self, aw, ah): self.width = aw; self.height = self._h(); return aw, self.height
    def draw(self):
        c = self.canv; aw = self.avail
        c.setFillColor(BG_SURFACE); c.roundRect(0, 0, aw, self.height, 6, fill=1, stroke=0)
        c.setStrokeColor(TEAL_DIM); c.setLineWidth(0.8); c.roundRect(0, 0, aw, self.height, 6, fill=0, stroke=1)
        if self.title:
            c.setFont("Inter-Bold", 8); c.setFillColor(TEAL); c.drawString(0.18*inch, self.height-0.21*inch, self.title)
            y = self.height-0.36*inch
        else:
            y = self.height-0.2*inch
        for it in self.items:
            ih = measure_wrapped_bold(it, 9, aw-0.55*inch); y -= ih; fy = y+ih-9*0.25
            c.setFont("Inter-Bold", 9); c.setFillColor(TEAL_DIM); c.drawString(0.18*inch, fy, "✓")
            draw_wrapped_bold(c, it, 0.36*inch, fy, 9, aw-0.52*inch, TEXT_GRAY, TEXT_WHITE); y -= 5

class ClosingNote(Flowable):
    def __init__(self, text):
        super().__init__(); self.text = text; self.avail = W-2*PAD
    def wrap(self, aw, ah):
        self.width = aw; self.height = measure_wrapped_bold(self.text, 9, self.avail-0.5*inch) + 0.38*inch; return aw, self.height
    def draw(self):
        c = self.canv; aw = self.avail
        c.setFillColor(BG_CARD_ALT); c.roundRect(0, 0, aw, self.height, 6, fill=1, stroke=0)
        c.setFillColor(SUCCESS); c.rect(0, 0, 3, self.height, fill=1, stroke=0)
        draw_wrapped_bold(c, self.text, 0.18*inch, self.height-0.22*inch, 9, aw-0.3*inch, TEXT_GRAY, TEXT_WHITE)

class FurtherReadingBox(Flowable):
    def __init__(self, items):
        super().__init__(); self.items = items; self.avail = W-2*PAD
    def _h(self):
        h = 0.32*inch
        for it in self.items: h += measure_wrapped_bold(it, 9, self.avail-0.55*inch) + 5
        return h + 0.12*inch
    def wrap(self, aw, ah): self.width = aw; self.height = self._h(); return aw, self.height
    def draw(self):
        c = self.canv; aw = self.avail
        c.setFillColor(BG_CARD); c.roundRect(0, 0, aw, self.height, 6, fill=1, stroke=0)
        c.setStrokeColor(BLUE_LIGHT); c.setLineWidth(0.8); c.roundRect(0, 0, aw, self.height, 6, fill=0, stroke=1)
        y = self.height-0.2*inch
        for it in self.items:
            ih = measure_wrapped_bold(it, 9, aw-0.55*inch); y -= ih; fy = y+ih-9*0.25
            c.setFont("Inter", 9); c.setFillColor(BLUE_LIGHT); c.drawString(0.18*inch, fy, "→")
            draw_wrapped_bold(c, it, 0.34*inch, fy, 9, aw-0.48*inch, BLUE_LIGHT, TEXT_WHITE); y -= 5

class RefWarningBox(Flowable):
    def __init__(self, text):
        super().__init__(); self.text = text; self.avail = W-2*PAD
    def wrap(self, aw, ah):
        self.width = aw
        self.height = measure_wrapped_bold(self.text, 8.7, self.avail-0.36*inch) + 0.35*inch
        return aw, self.height
    def draw(self):
        c = self.canv; aw = self.avail
        c.setFillColor(BG_SURFACE); c.roundRect(0, 0, aw, self.height, 5, fill=1, stroke=0)
        c.setStrokeColor(WARN); c.setLineWidth(0.8); c.roundRect(0, 0, aw, self.height, 5, fill=0, stroke=1)
        draw_wrapped_bold(c, self.text, 0.18*inch, self.height-0.21*inch, 8.7, aw-0.3*inch, WARN, TEXT_WHITE)

class ImageRefBlock(Flowable):
    def __init__(self, image_path, max_height, size_mode="full", crop_start=0.0, crop_end=1.0):
        super().__init__()
        self.image_path = image_path
        self.max_height = max_height
        self.size_mode = size_mode if size_mode in IMAGE_SIZE_OPTIONS else "full"
        self.crop_start = max(0.0, min(1.0, crop_start))
        self.crop_end = max(self.crop_start, min(1.0, crop_end))
        self._image_reader = None
        self._draw_w = 1.0
        self._draw_h_total = 1.0
        self._draw_h_part = 1.0
        self._vpad = 0.08 * inch

    def _compute_size(self, avail_w):
        if self._image_reader is None:
            self._image_reader = ImageReader(self.image_path)
        img_w, img_h = self._image_reader.getSize()
        max_h = max(1.0, self.max_height)
        self._draw_w, self._draw_h_total = _compute_image_size(
            img_w,
            img_h,
            max(1.0, avail_w),
            max_h,
            self.size_mode,
        )
        self._draw_h_part = max(1.0, self._draw_h_total * (self.crop_end - self.crop_start))
        if self._draw_h_part > max_h:
            # Cota de seguridad para no romper paginación si un split manual aún resulta alto.
            shrink = max_h / self._draw_h_part
            self._draw_w *= shrink
            self._draw_h_total *= shrink
            self._draw_h_part = max_h

    def wrap(self, aw, ah):
        self.width = aw
        self._compute_size(aw)
        self.height = self._draw_h_part + (2 * self._vpad)
        return aw, self.height

    def split(self, aw, ah):
        return []

    def _draw_shadow(self, c, x, y, w, h):
        if not hasattr(c, "setFillAlpha"):
            return
        c.saveState()
        c.setFillAlpha(IMAGE_SHADOW_ALPHA)
        c.setFillColor(HexColor("#000000"))
        c.roundRect(
            x + IMAGE_SHADOW_OFFSET,
            y - IMAGE_SHADOW_OFFSET,
            w,
            h,
            2,
            fill=1,
            stroke=0,
        )
        c.restoreState()

    def _draw_border(self, c, x, y, w, h):
        c.saveState()
        c.setStrokeColor(IMAGE_FRAME_COLOR)
        c.setLineWidth(IMAGE_FRAME_LINE_WIDTH)
        c.roundRect(x, y, w, h, 2, fill=0, stroke=1)
        c.restoreState()

    def draw(self):
        x = max(0.0, (self.width - self._draw_w) / 2.0)
        y = self._vpad
        c = self.canv
        self._draw_shadow(c, x, y, self._draw_w, self._draw_h_part)
        if self.crop_start <= 0.0 and self.crop_end >= 1.0:
            c.drawImage(self.image_path, x, y, width=self._draw_w, height=self._draw_h_total, mask="auto")
            self._draw_border(c, x, y, self._draw_w, self._draw_h_part)
            return
        c.saveState()
        path = c.beginPath()
        path.rect(x, y, self._draw_w, self._draw_h_part)
        c.clipPath(path, stroke=0, fill=0)
        # PDF usa origen abajo-izquierda; para que el split siga orden top->bottom
        # (igual que DOCX), el anclaje debe desplazarse con crop_end.
        offset_y = self._draw_h_total * (1.0 - self.crop_end)
        c.drawImage(
            self.image_path,
            x,
            y - offset_y,
            width=self._draw_w,
            height=self._draw_h_total,
            mask="auto",
        )
        c.restoreState()
        self._draw_border(c, x, y, self._draw_w, self._draw_h_part)

SP = lambda n=1: Spacer(1, n*0.13*inch)

# ── Parser de Markdown (patrones en PORTUGUÉS) ───────────────────────────────
_IMPERATIVE = [
    'abra ', 'faça ', 'revise ', 'selecione ', 'verifique ', 'confirme ',
    'salve ', 'guarde ', 'execute ', 'lance ', 'localize ', 'insira ',
    'atribua ', 'adicione ', 'aguarde ', 'volte ', 'decida ', 'identifique ',
    'detecte ', 'relacione ', 'evite ', 'documente ', 'conserve ', 'mantenha ',
    'duplique ', 'crie ', 'marque ', 'repita ', 'defina ', 'use ', 'ajuste ',
    'retome ', 'atualize ', 'filtre ', 'acesse ', 'navegue ', 'clique ',
    'carregue ', 'baixe ', 'importe ', 'exporte ', 'pressione ', 'configure ',
    'no goalbus', 'dentro do', 'dentro da', 'se o ', 'se a ', 'se não ', 'se exist',
    'na barra', 'no módulo', 'na tela', 'em goalbus', 'na seção',
    'abra o ', 'abra a ', 'vá para', 'vá ao ', 'vá à ',
]

def _is_instruction(items):
    if not items: return False
    first = items[0].get('text', '').lower().strip()
    return any(first.startswith(v) for v in _IMPERATIVE)

def _classify(items, ctx, is_bullet=False):
    c = ctx.lower()
    if any(k in c for k in ['antes de começar', 'antes de continuar', 'antes de prosseguir',
                              'antes de terminar', 'certifique-se de que', 'verifique se você']):
        return 'prereq'
    if 'para o caso de referência' in c or 'para o seu caso de referência' in c:
        is_q = any(k in c for k in ['confirme', 'certifique-se', 'afirmar', 'não continue',
                                     'não prossiga', 'não considere', 'termine este quick',
                                     'apenas quando puder', 'pergunte-se', 'pergunte se'])
        if is_q or (items and items[0].get('text', '').startswith('?')):
            return 'question'
        return 'example_bullet' if is_bullet else 'example_num'
    if is_bullet:
        return 'further'
    if _is_instruction(items):
        return 'steps'
    return 'inline'

def _parse_list(lines, i):
    items = []
    while i < len(lines):
        line = lines[i]
        m = re.match(r'^(\d+)\.\s+(.*)', line)
        if m: items.append({'snum': int(m.group(1)), 'text': m.group(2).strip()}); i += 1; continue
        m = re.match(r'^[-*+]\s+(.*)', line)
        if m: items.append({'snum': '•', 'text': m.group(1).strip()}); i += 1; continue
        m = re.match(r'^\s{2,}(\d+)\.\s+(.*)', line)
        if m:
            if items:
                items[-1].setdefault('subs', []).append({'snum': int(m.group(1)), 'text': m.group(2).strip()})
            i += 1; continue
        m = re.match(r'^\s{2,}[-*+]\s+(.*)', line)
        if m:
            if items:
                items[-1].setdefault('subs', []).append({'snum': '•', 'text': m.group(1).strip()})
            i += 1; continue
        if not line.strip():
            j = i+1
            while j < len(lines) and not lines[j].strip(): j += 1
            if j < len(lines) and (re.match(r'^\d+\.', lines[j]) or re.match(r'^[-*+]\s', lines[j]) or re.match(r'^\s{2,}', lines[j])):
                i = j; continue
            else: break
        break
    return items, i

def build_pdf_table(rows, avail_w):
    if not rows: return Spacer(1, 1)
    data = []
    for r in rows:
        data.append([Paragraph(cell, ParagraphStyle('tc', fontName='Inter', fontSize=8.5, textColor=TEXT_WHITE, leading=10)) for cell in r])
    
    col_count = len(rows[0])
    col_w = avail_w / col_count
    
    # Auto-escalar fuente si hay demasiadas columnas
    fsize = 8.5
    if col_count > 6: fsize = 7
    if col_count > 9: fsize = 5.5
    
    for r_idx in range(len(data)):
        for c_idx in range(len(data[r_idx])):
             data[r_idx][c_idx].style.fontSize = fsize
             if r_idx == 0: data[r_idx][c_idx].style.fontName = 'Inter-Bold'
    
    t = Table(data, colWidths=[col_w]*col_count)
    t.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), BG_CARD_ALT),
        ('BACKGROUND', (0,1), (-1,-1), BG_CARD),
        ('GRID', (0,0), (-1,-1), 0.5, BG_SURFACE),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('LEFTPADDING', (0,0), (-1,-1), 6),
        ('RIGHTPADDING', (0,0), (-1,-1), 6),
        ('TOPPADDING', (0,0), (-1,-1), 4),
        ('BOTTOMPADDING', (0,0), (-1,-1), 4),
    ]))
    return t

def extract_ref_name(line):
    m = REF_LINE_RE.match(line or "")
    if not m:
        return None
    return m.group(1).strip()

def parse_ref_directive(ref_text):
    raw = (ref_text or "").strip()
    parts = [p.strip() for p in raw.split("|")]
    name = parts[0] if parts else ""
    spec = {
        "raw": raw,
        "name": name,
        "size": DEFAULT_IMAGE_SIZE,
        "split": None,
        "warnings": [],
    }
    for opt in parts[1:]:
        if not opt:
            continue
        opt_l = opt.lower()
        # Atajo amigable: "ref: img | compact" o "ref: img | 2"
        if "=" not in opt:
            if opt_l in IMAGE_SIZE_OPTIONS:
                spec["size"] = opt_l
                continue
            if opt_l.isdigit():
                n = int(opt_l)
                if 1 <= n <= MAX_IMAGE_SPLIT_PARTS:
                    spec["split"] = n
                else:
                    spec["warnings"].append(f"split fuera de rango '{opt}', usa 1..{MAX_IMAGE_SPLIT_PARTS}")
                continue
            spec["warnings"].append(f"opcion invalida '{opt}' (usa clave=valor, p.ej. size=compact)")
            continue
        key, value = [x.strip() for x in opt.split("=", 1)]
        key = key.lower()
        value_l = value.lower()
        if key == "size":
            if value_l in IMAGE_SIZE_OPTIONS:
                spec["size"] = value_l
            else:
                spec["warnings"].append(f"size invalido '{value}', usa: {', '.join(IMAGE_SIZE_OPTIONS)}")
        elif key == "split":
            try:
                n = int(value)
                if 1 <= n <= MAX_IMAGE_SPLIT_PARTS:
                    spec["split"] = n
                else:
                    spec["warnings"].append(f"split fuera de rango '{value}', usa 1..{MAX_IMAGE_SPLIT_PARTS}")
            except Exception:
                spec["warnings"].append(f"split invalido '{value}', usa entero")
        else:
            spec["warnings"].append(f"opcion desconocida '{key}'")
    return spec

def _auto_image_size_mode(img_w, img_h):
    ratio = img_w / float(max(1, img_h))
    if ratio >= 1.35:
        return "full"
    if ratio <= 0.95:
        return "compact"
    return "full" if img_w >= 1400 else "compact"

def _resolve_image_mode(spec, img_w, img_h):
    chosen = (spec or {}).get("size", DEFAULT_IMAGE_SIZE)
    if chosen in ("full", "compact"):
        return chosen
    return _auto_image_size_mode(img_w, img_h)

def _compact_height_cap(max_h, img_w, img_h):
    ratio = img_w / float(max(1, img_h))
    if ratio <= 0.95:
        frac = DEFAULT_COMPACT_HEIGHT_FRAC_VERTICAL
    elif ratio >= 1.6:
        frac = DEFAULT_COMPACT_HEIGHT_FRAC_PANORAMIC
    else:
        frac = DEFAULT_COMPACT_HEIGHT_FRAC
    return max(1.0, max_h * frac)

def _compute_image_size(img_w, img_h, avail_w, max_h, mode):
    if mode == "full":
        target_w = avail_w
    else:
        target_w = avail_w * DEFAULT_COMPACT_WIDTH_RATIO

    scale = min(1.0, target_w / float(max(1, img_w)))
    draw_w = max(1.0, img_w * scale)
    draw_h = max(1.0, img_h * scale)

    if mode == "compact":
        compact_h_cap = _compact_height_cap(max_h, img_w, img_h)
        if draw_h > compact_h_cap:
            shrink = compact_h_cap / draw_h
            draw_w *= shrink
            draw_h = compact_h_cap

    return draw_w, draw_h

def _build_image_plan(image_path, spec, avail_w, max_h):
    img_w, img_h = ImageReader(image_path).getSize()
    requested_mode = (spec or {}).get("size", DEFAULT_IMAGE_SIZE)
    mode = _resolve_image_mode(spec, img_w, img_h)
    draw_w, draw_h = _compute_image_size(img_w, img_h, avail_w, max_h, mode)

    # En auto, una panorámica detallada puede subir a full para evitar ilegibilidad.
    if (
        mode == "compact"
        and requested_mode == "auto"
        and img_w >= 1400
        and (img_w / float(max(1, img_h))) >= 1.2
        and draw_h < (max_h * 0.45)
    ):
        mode = "full"
        draw_w, draw_h = _compute_image_size(img_w, img_h, avail_w, max_h, mode)

    split = (spec or {}).get("split")
    if split and split > 1:
        parts = split
    else:
        parts = int(math.ceil(draw_h / float(max(1.0, max_h)))) if draw_h > max_h else 1
    parts = max(1, min(MAX_IMAGE_SPLIT_PARTS, parts))

    if parts == 1:
        ranges = [(0.0, 1.0)]
    else:
        step = 1.0 / parts
        ranges = [(i * step, (i + 1) * step) for i in range(parts)]

    return {
        "mode": mode,
        "draw_w": draw_w,
        "draw_h": draw_h,
        "parts": parts,
        "ranges": ranges,
        "src_w": img_w,
        "src_h": img_h,
    }

def _crop_image_to_stream(image_path, start_frac, end_frac):
    if PILImage is None:
        return None
    start_frac = max(0.0, min(1.0, start_frac))
    end_frac = max(start_frac, min(1.0, end_frac))
    with PILImage.open(image_path) as im:
        width, height = im.size
        top = int(round(height * start_frac))
        bottom = int(round(height * end_frac))
        if bottom <= top:
            bottom = min(height, top + 1)
        cut = im.crop((0, top, width, bottom))
        fmt = (im.format or "PNG").upper()
        if fmt not in ("PNG", "JPEG"):
            fmt = "PNG"
        if fmt == "JPEG" and cut.mode not in ("RGB", "L"):
            cut = cut.convert("RGB")
        stream = io.BytesIO()
        save_kwargs = {"quality": 95, "subsampling": 0} if fmt == "JPEG" else {}
        cut.save(stream, format=fmt, **save_kwargs)
        stream.seek(0)
        return stream

def _guide_dir_candidates(md_dir, p_num):
    dirs = [os.path.join(md_dir, f"P{p_num}")]
    p_padded = f"P{p_num:02d}"
    if p_padded != f"P{p_num}":
        dirs.append(os.path.join(md_dir, p_padded))
    return dirs

def _candidate_ref_names(ref_raw):
    name = (ref_raw or "").strip().strip("'\"")
    if not name or "/" in name or "\\" in name:
        return []
    base, ext = os.path.splitext(name)
    if ext:
        if ext.lower() in ALLOWED_IMAGE_EXTS:
            return [name]
        # Soporta nombres sin extension que incluyen puntos (ej. "...11.27.55").
        if ext[1:].isalpha() and len(ext) <= 5:
            return []
        return [f"{name}{extn}" for extn in ALLOWED_IMAGE_EXTS]
    return [f"{base}{extn}" for extn in ALLOWED_IMAGE_EXTS]

def _find_file_case_insensitive(folder, file_name):
    if not os.path.isdir(folder):
        return None
    target = file_name.lower()
    for item in os.listdir(folder):
        if item.lower() == target:
            path = os.path.join(folder, item)
            if os.path.isfile(path):
                return path
    return None

def resolve_ref_image(md_dir, p_num, ref_raw):
    names = _candidate_ref_names(ref_raw)
    if not names:
        return None
    for folder in _guide_dir_candidates(md_dir, p_num):
        for name in names:
            candidate = os.path.join(folder, name)
            if os.path.isfile(candidate):
                return candidate
            candidate_ci = _find_file_case_insensitive(folder, name)
            if candidate_ci:
                return candidate_ci
    return None

def _validate_image_refs(blocks, md_dir, p_num, ref_cache, log_fn):
    searched_dirs = [os.path.basename(d) for d in _guide_dir_candidates(md_dir, p_num)]
    for block in blocks:
        if block[0] == 'image_ref_invalid':
            log_fn(f"    [ref] AVISO P{p_num}: referencia invalida ({block[1] or 'ref:'})")
            continue
        if block[0] != 'image_ref':
            continue
        ref_spec = block[1]
        for warn in ref_spec.get("warnings", []):
            log_fn(f"    [ref] AVISO P{p_num}: {warn} en 'ref: {ref_spec.get('raw', '')}'")
        ref_name = ref_spec.get("name", "")
        key = (p_num, ref_name)
        if key in ref_cache:
            continue
        image_path = resolve_ref_image(md_dir, p_num, ref_name)
        ref_cache[key] = image_path
        if image_path:
            rel = os.path.relpath(image_path, md_dir)
            opts = []
            if ref_spec.get("size", DEFAULT_IMAGE_SIZE) != DEFAULT_IMAGE_SIZE:
                opts.append(f"size={ref_spec.get('size')}")
            if ref_spec.get("split"):
                opts.append(f"split={ref_spec.get('split')}")
            opts_txt = f" ({', '.join(opts)})" if opts else ""
            log_fn(f"    [ref] OK P{p_num}: '{ref_name}' -> {rel}{opts_txt}")
        else:
            log_fn(f"    [ref] AVISO P{p_num}: no se encontro '{ref_name}' en {', '.join(searched_dirs)}")

def _parse_section_body(body):
    lines = body.split('\n')
    blocks = []
    pending = []

    def flush():
        if pending:
            t = ' '.join(pending).strip()
            if t:
                if any(k in t.lower() for k in ['quando terminar', 'quando concluir',
                                                  'ao terminar', 'ao concluir']):
                    blocks.append(('closing', t))
                else:
                    blocks.append(('body', t))
            pending.clear()

    i = 0
    while i < len(lines):
        line = lines[i]
        if not line.strip(): flush(); i += 1; continue
        
        # Detección de tablas
        if line.strip().startswith('|'):
            flush()
            t_rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                l = lines[i].strip()
                if not re.match(r'^\|[\s\-:|]+\|$', l): # Ignorar linea separadora Markdown
                    # Limpiar los separadores | de los extremos antes de hacer el split
                    l_clean = l.strip('|').strip()
                    cells = [c.strip() for c in l_clean.split('|')]
                    if cells: t_rows.append(cells)
                i += 1
            if t_rows: blocks.append(('table', t_rows))
            continue
            
        ref_name = extract_ref_name(line)
        if ref_name is not None:
            flush()
            if ref_name:
                ref_spec = parse_ref_directive(ref_name)
                if ref_spec.get("name"):
                    blocks.append(('image_ref', ref_spec))
                else:
                    blocks.append(('image_ref_invalid', line.strip()))
            else:
                blocks.append(('image_ref_invalid', line.strip()))
            i += 1
            continue
        if line.startswith('>'):
            flush()
            bq = []
            while i < len(lines) and lines[i].startswith('>'):
                bq.append(lines[i].lstrip('> ').strip()); i += 1
            text = ' '.join(bq)
            blocks.append(('case_ref', text)); continue
        if re.match(r'^\d+\.\s', line):
            ctx = ' '.join(pending); flush()
            items, i = _parse_list(lines, i)
            kind = _classify(items, ctx, False)
            if kind == 'prereq':
                blocks.append(('prereq', [it['text'] for it in items]))
            elif kind == 'steps':
                blocks.append(('steps', items))
            elif kind == 'question':
                blocks.append(('question', '', [it['text'] for it in items]))
            elif kind == 'example_num':
                blocks.append(('example', '', [it['text'] for it in items], True))
            else:
                blocks.append(('steps', items))
            continue
        if re.match(r'^[-*+]\s', line):
            ctx = ' '.join(pending); flush()
            items, i = _parse_list(lines, i)
            kind = _classify(items, ctx, True)
            link_items = []
            for it in items:
                t = it['text']
                lm = re.match(r'\[([^\]]+)\]\([^\)]+\)', t)
                link_items.append(lm.group(1) if lm else t)
            if kind == 'further':
                blocks.append(('further', link_items))
            elif kind in ('example_bullet', 'example_num'):
                blocks.append(('example', '', link_items, False))
            elif kind == 'question':
                blocks.append(('question', '', link_items))
            else:
                blocks.append(('steps', items))
            continue
        pending.append(line.strip()); i += 1

    flush()
    return blocks

def _blocks_to_flowables(blocks, sec_title, sec_num, is_further, md_dir, p_num, ref_cache, image_max_h):
    fl = [KeepTogether([SectionHeader(sec_title, sec_num if not is_further else None), SP()])]
    for b in blocks:
        kind = b[0]
        if kind == 'body':       fl += [BodyText(b[1]), SP()]
        elif kind == 'table':     fl += [build_pdf_table(b[1], W-2*PAD), SP()]
        elif kind == 'closing':  fl += [ClosingNote(b[1]), SP(2)]
        elif kind == 'case_ref': fl += [CaseRefBox(b[1]), SP()]
        elif kind == 'prereq':   fl += [PrereqBox(b[1]), SP()]
        elif kind == 'steps':    fl += [NumberedStepsList(b[1]), SP()]
        elif kind == 'question': fl += [QuestionList(b[1].strip().rstrip(':'), b[2]), SP()]
        elif kind == 'example':  fl += [ExampleBox(b[1].strip().rstrip(':'), b[2], numbered=b[3]), SP()]
        elif kind == 'further':  fl += [FurtherReadingBox(b[1]), SP()]
        elif kind == 'image_ref':
            ref_spec = b[1]
            ref_name = ref_spec.get("name", "")
            image_path = ref_cache.get((p_num, ref_name))
            if image_path:
                plan = _build_image_plan(image_path, ref_spec, W - 2 * PAD, image_max_h)
                for idx, (crop_start, crop_end) in enumerate(plan["ranges"]):
                    fl += [KeepTogether([
                        ImageRefBlock(
                            image_path,
                            image_max_h,
                            size_mode=plan["mode"],
                            crop_start=crop_start,
                            crop_end=crop_end,
                        ),
                        SP(0.5 if idx < len(plan["ranges"]) - 1 else 1),
                    ])]
            else:
                fl += [RefWarningBox(f"Imagen no encontrada para referencia: ref: {ref_name}"), SP()]
        elif kind == 'image_ref_invalid':
            fl += [RefWarningBox("Referencia de imagen vacia o invalida. Usa: ref: nombre_imagen"), SP()]
    return fl

# ── Builder principal ────────────────────────────────────────────────────────
def build_pdf(md_path, out_dir, log_fn=print):
    basename = os.path.basename(md_path)
    base_no_ext = os.path.splitext(basename)[0]
    m = re.match(r'^P(\d+)', basename, re.IGNORECASE)
    p_num = int(m.group(1)) if m else 0
    md_dir = os.path.dirname(os.path.abspath(md_path))

    with open(md_path, 'r', encoding='utf-8') as f:
        raw = f.read()
    raw = re.sub(r'\s*filecite\S*', '', raw)

    meta = {}
    fm = re.match(r'^---\s*\n(.*?)\n---\s*\n', raw, re.DOTALL)
    if fm:
        for line in fm.group(1).split('\n'):
            mm = re.match(r'^(\w+):\s*(.*)', line)
            if mm: meta[mm.group(1)] = mm.group(2).strip().strip("'\"")
        raw = raw[fm.end():]
    title = meta.get('title', f'Quick Start P{p_num}')
    intro = meta.get('intro', '')

    parts = re.split(r'\n## ', '\n'+raw)
    sections = [(p.split('\n', 1)[0].strip(), p.split('\n', 1)[1] if '\n' in p else '') for p in parts[1:]]

    bg = make_bg(f"goalbus  •  Quick Start P{p_num}", f"goalbus  •  {title[:60]}")
    story = [HeroBlock(p_num, title, intro), SP(2)]
    image_max_h = (H - TOP_MARGIN - BOTTOM_MARGIN) * CONTENT_MAX_IMAGE_FRAC

    parsed_sections = []
    sec_num = 0
    for sec_title, sec_body in sections:
        sec_title = re.sub(r'^\d+[\.\-]?\s+', '', sec_title)
        is_further = any(k in sec_title.lower() for k in ['leituras adicionais', 'lecturas adicionales'])
        if not is_further:
            sec_num += 1
        blocks = _parse_section_body(sec_body)
        parsed_sections.append((sec_title, sec_num if not is_further else None, is_further, blocks))

    ref_cache = {}
    for _, _, _, blocks in parsed_sections:
        _validate_image_refs(blocks, md_dir, p_num, ref_cache, log_fn)

    for sec_title, sec_index, is_further, blocks in parsed_sections:
        story.extend(_blocks_to_flowables(
            blocks,
            sec_title,
            sec_index,
            is_further,
            md_dir,
            p_num,
            ref_cache,
            image_max_h,
        ))

    out_path = os.path.join(out_dir, f"{base_no_ext}.pdf")
    doc = SimpleDocTemplate(out_path, pagesize=letter,
                            leftMargin=PAD, rightMargin=PAD,
                            topMargin=TOP_MARGIN, bottomMargin=BOTTOM_MARGIN)
    doc.build(story, onFirstPage=bg, onLaterPages=bg)
    return out_path

def _collect_document_structure(md_path, log_fn):
    basename = os.path.basename(md_path)
    m = re.match(r'^P(\d+)', basename, re.IGNORECASE)
    p_num = int(m.group(1)) if m else 0
    md_dir = os.path.dirname(os.path.abspath(md_path))

    with open(md_path, 'r', encoding='utf-8') as f:
        raw = f.read()
    raw = re.sub(r'\s*filecite\S*', '', raw)

    meta = {}
    fm = re.match(r'^---\s*\n(.*?)\n---\s*\n', raw, re.DOTALL)
    if fm:
        for line in fm.group(1).split('\n'):
            mm = re.match(r'^(\w+):\s*(.*)', line)
            if mm:
                meta[mm.group(1)] = mm.group(2).strip().strip("'\"")
        raw = raw[fm.end():]

    title = meta.get('title', f'Quick Start P{p_num}')
    intro = meta.get('intro', '')

    parts = re.split(r'\n## ', '\n' + raw)
    sections = [
        (p.split('\n', 1)[0].strip(), p.split('\n', 1)[1] if '\n' in p else '')
        for p in parts[1:]
    ]

    parsed_sections = []
    sec_num = 0
    for sec_title, sec_body in sections:
        sec_title = re.sub(r'^\d+[\.\-]?\s+', '', sec_title)
        is_further = any(k in sec_title.lower() for k in ['leituras adicionais', 'lecturas adicionales'])
        if not is_further:
            sec_num += 1
        blocks = _parse_section_body(sec_body)
        parsed_sections.append((sec_title, sec_num if not is_further else None, is_further, blocks))

    ref_cache = {}
    for _, _, _, blocks in parsed_sections:
        _validate_image_refs(blocks, md_dir, p_num, ref_cache, log_fn)

    return {
        "p_num": p_num,
        "md_dir": md_dir,
        "title": title,
        "intro": intro,
        "parsed_sections": parsed_sections,
        "ref_cache": ref_cache,
    }

def _docx_set_page_background(document, fill_hex, OxmlElement, qn):
    # Eliminar fondos existentes
    for node in document.element.findall(qn("w:background")):
        document.element.remove(node)
    
    # Crear el nodo background con el color oscuro
    bg = OxmlElement("w:background")
    bg.set(qn("w:color"), fill_hex)
    document.element.insert(0, bg)
    
    # FORZAR VISIBILIDAD: Activar visualización de formas de fondo en los ajustes del documento
    settings = document.settings.element
    display_bg = settings.find(qn("w:displayBackgroundShape"))
    if display_bg is None:
        display_bg = OxmlElement("w:displayBackgroundShape")
        settings.append(display_bg)
    display_bg.set(qn("w:val"), "true")

def _docx_set_cell_background(cell, fill_hex, OxmlElement, qn):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)

def _docx_set_cell_borders(cell, color_hex, OxmlElement, qn, size=8):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        edge_tag = qn(f"w:{edge}")
        element = borders.find(edge_tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color_hex)

def _docx_set_cell_margins(cell, OxmlElement, qn, top=110, right=120, bottom=110, left=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "right": right, "bottom": bottom, "left": left}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

def _docx_clear_table_borders(table, OxmlElement, qn):
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_tag = qn(f"w:{edge}")
        node = tbl_borders.find(edge_tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tbl_borders.append(node)
        node.set(qn("w:val"), "nil")

def _docx_add_page_field(paragraph, OxmlElement, qn):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_sep)
    run._r.append(fld_end)

def _docx_style_inline_picture(inline_shape, OxmlElement):
    sp_pr_nodes = inline_shape._inline.xpath(".//pic:spPr")
    if not sp_pr_nodes:
        return
    sp_pr = sp_pr_nodes[0]

    for child in list(sp_pr):
        if child.tag.endswith("}ln") or child.tag.endswith("}effectLst"):
            sp_pr.remove(child)

    line = OxmlElement("a:ln")
    line.set("w", DOCX_IMAGE_FRAME_WIDTH)
    solid_fill = OxmlElement("a:solidFill")
    frame_color = OxmlElement("a:srgbClr")
    frame_color.set("val", DOCX_IMAGE_FRAME)
    solid_fill.append(frame_color)
    line.append(solid_fill)
    dash = OxmlElement("a:prstDash")
    dash.set("val", "solid")
    line.append(dash)
    sp_pr.append(line)

    effect_list = OxmlElement("a:effectLst")
    outer_shadow = OxmlElement("a:outerShdw")
    outer_shadow.set("blurRad", DOCX_IMAGE_SHADOW_BLUR)
    outer_shadow.set("dist", DOCX_IMAGE_SHADOW_DIST)
    outer_shadow.set("dir", "5400000")
    outer_shadow.set("algn", "ctr")
    outer_shadow.set("rotWithShape", "0")
    shadow_color = OxmlElement("a:srgbClr")
    shadow_color.set("val", "000000")
    shadow_alpha = OxmlElement("a:alpha")
    shadow_alpha.set("val", DOCX_IMAGE_SHADOW_ALPHA)
    shadow_color.append(shadow_alpha)
    outer_shadow.append(shadow_color)
    effect_list.append(outer_shadow)
    sp_pr.append(effect_list)

def _docx_add_bold_runs(paragraph, text, Pt, RGBColor, color_hex, bold_color_hex=None, size_pt=10, force_bold=False):
    bold_color = bold_color_hex or color_hex
    for segment, is_bold in parse_bold(text):
        if not segment:
            continue
        run = paragraph.add_run(segment)
        run.bold = force_bold or is_bold
        run.font.name = "Arial"
        run.font.size = Pt(size_pt)
        run.font.color.rgb = RGBColor.from_string(bold_color if (force_bold or is_bold) else color_hex)

def _docx_style_paragraph(paragraph, Pt, Inches, align=None, after=4, line=1.22, left_indent=0.0, right_indent=0.0):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if left_indent > 0:
        pf.left_indent = Inches(left_indent)
    if right_indent > 0:
        pf.right_indent = Inches(right_indent)
    if align is not None:
        paragraph.alignment = align

def _docx_circled_marker(marker):
    marker_txt = str(marker).strip()
    if marker_txt.isdigit():
        n = int(marker_txt)
        if 1 <= n <= 20:
            return chr(0x2460 + n - 1)  # ①..⑳
    if marker_txt == "•":
        return "●"
    return marker_txt

def _docx_add_card(document, Inches, WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT, OxmlElement, qn,
                   content_w_in, bg_hex, border_hex=None, accent_hex=None):
    cols = 2 if accent_hex else 1
    table = document.add_table(rows=1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    accent_w_in = 0.05

    if accent_hex:
        left_cell = table.cell(0, 0)
        main_cell = table.cell(0, 1)
        left_cell.width = Inches(accent_w_in)
        main_cell.width = Inches(max(0.5, content_w_in - accent_w_in))
        _docx_set_cell_background(left_cell, accent_hex, OxmlElement, qn)
        _docx_set_cell_margins(left_cell, OxmlElement, qn, top=35, right=35, bottom=35, left=35)
        left_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    else:
        main_cell = table.cell(0, 0)
        main_cell.width = Inches(content_w_in)

    _docx_set_cell_background(main_cell, bg_hex, OxmlElement, qn)
    _docx_set_cell_margins(main_cell, OxmlElement, qn, top=110, right=140, bottom=110, left=140)
    if border_hex:
        _docx_set_cell_borders(main_cell, border_hex, OxmlElement, qn, size=8)
    main_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    return main_cell

def _docx_add_badge_row(container_cell, marker, text, text_color_hex, bold_color_hex, marker_bg_hex, marker_fg_hex,
                        item_bg_hex,
                        OxmlElement, qn, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH,
                        badge_w_in=0.24, text_w_in=5.0, text_size_pt=9):
    row_table = container_cell.add_table(rows=1, cols=2)
    row_table.autofit = False
    _docx_clear_table_borders(row_table, OxmlElement, qn)

    marker_cell = row_table.cell(0, 0)
    text_cell = row_table.cell(0, 1)
    marker_cell.width = Inches(badge_w_in)
    text_cell.width = Inches(max(0.5, text_w_in - badge_w_in))

    _docx_set_cell_background(marker_cell, item_bg_hex, OxmlElement, qn)
    _docx_set_cell_background(text_cell, item_bg_hex, OxmlElement, qn)
    _docx_set_cell_margins(marker_cell, OxmlElement, qn, top=35, right=45, bottom=35, left=45)
    _docx_set_cell_margins(text_cell, OxmlElement, qn, top=0, right=45, bottom=0, left=95)

    marker_p = marker_cell.paragraphs[0]
    _docx_style_paragraph(marker_p, Pt, Inches, align=WD_ALIGN_PARAGRAPH.CENTER, after=0, line=1.0)
    marker_run = marker_p.add_run(_docx_circled_marker(marker))
    marker_run.bold = True
    marker_run.font.name = "Arial"
    marker_run.font.size = Pt(9.2)
    marker_run.font.color.rgb = RGBColor.from_string(marker_bg_hex or marker_fg_hex)

    text_cell.text = ""
    text_p = text_cell.paragraphs[0]
    _docx_style_paragraph(text_p, Pt, Inches, after=1.6, line=1.2, right_indent=0.12)
    _docx_add_bold_runs(text_p, text, Pt, RGBColor, text_color_hex, bold_color_hex=bold_color_hex, size_pt=text_size_pt)
    return text_cell

def build_docx(md_path, out_dir, log_fn=print):
    try:
        from docx import Document
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:
        raise RuntimeError("Falta 'python-docx'. Instala con: pip install python-docx") from exc

    context = _collect_document_structure(md_path, log_fn)
    p_num = context["p_num"]
    md_dir = context["md_dir"]
    title = context["title"]
    intro = context["intro"]
    parsed_sections = context["parsed_sections"]
    ref_cache = context["ref_cache"]

    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(PAD / inch)
    section.right_margin = Inches(PAD / inch)
    section.top_margin = Inches(TOP_MARGIN / inch)
    section.bottom_margin = Inches(BOTTOM_MARGIN / inch)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(DOCX_TEXT_GRAY)

    _docx_set_page_background(document, DOCX_BG_PAGE, OxmlElement, qn)

    # Header
    header = section.header
    for p in header.paragraphs:
        p.text = ""
    hp_logo = header.paragraphs[0]
    hp_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if LOGO_PATH and os.path.exists(LOGO_PATH):
        # Auto-switch a gris opaco para contraste sobre fondo de hoja blanca
        logo_cand = os.path.join(md_dir, 'goal-logo-anthracite.png')
        if not os.path.exists(logo_cand): logo_cand = os.path.join(os.path.dirname(md_dir), 'goal-logo-anthracite.png')
        final_logo = logo_cand if os.path.exists(logo_cand) else LOGO_PATH
        hp_logo.add_run().add_picture(final_logo, width=Inches(1.1))
    _docx_style_paragraph(hp_logo, Pt, Inches, after=2, line=1.0)

    hp_label = header.add_paragraph()
    hp_label.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _docx_style_paragraph(hp_label, Pt, Inches, after=0, line=1.0)
    r = hp_label.add_run(f"goalbus  •  Quick Start P{p_num}")
    r.font.name = "Arial"
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(DOCX_TEXT_DIM)

    # Footer
    footer = section.footer
    for p in footer.paragraphs:
        p.text = ""
    fp_left = footer.paragraphs[0]
    _docx_style_paragraph(fp_left, Pt, Inches, after=0, line=1.0)
    fp_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rl = fp_left.add_run(f"goalbus  •  {title[:60]}")
    rl.font.name = "Arial"
    rl.font.size = Pt(7.5)
    rl.font.color.rgb = RGBColor.from_string(DOCX_TEXT_DIM)

    fp_right = footer.add_paragraph()
    _docx_style_paragraph(fp_right, Pt, Inches, after=0, line=1.0)
    fp_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = fp_right.add_run("Página ")
    rr.font.name = "Arial"
    rr.font.size = Pt(7.5)
    rr.font.color.rgb = RGBColor.from_string(DOCX_TEXT_DIM)
    _docx_add_page_field(fp_right, OxmlElement, qn)

    content_w_in = (W - 2 * PAD) / inch
    image_max_h_in = ((H - TOP_MARGIN - BOTTOM_MARGIN) * CONTENT_MAX_IMAGE_FRAC) / inch

    # Hero
    hero_cell = _docx_add_card(
        document,
        Inches,
        WD_TABLE_ALIGNMENT,
        WD_CELL_VERTICAL_ALIGNMENT,
        OxmlElement,
        qn,
        content_w_in,
        DOCX_BG_CARD,
        border_hex=None,
        accent_hex=DOCX_TEAL_DIM,
    )
    hero_cell.text = ""
    p0 = hero_cell.paragraphs[0]
    _docx_style_paragraph(p0, Pt, Inches, after=4, line=1.0)
    _docx_add_bold_runs(p0, f"P{p_num}", Pt, RGBColor, DOCX_TEAL, size_pt=11, force_bold=True)

    p1 = hero_cell.add_paragraph()
    _docx_style_paragraph(p1, Pt, Inches, after=5, line=1.2)
    _docx_add_bold_runs(p1, title, Pt, RGBColor, DOCX_TEXT_WHITE, bold_color_hex=DOCX_TEAL, size_pt=18)

    p2 = hero_cell.add_paragraph()
    _docx_style_paragraph(p2, Pt, Inches, after=0, line=1.25)
    _docx_add_bold_runs(p2, intro, Pt, RGBColor, DOCX_TEXT_GRAY, bold_color_hex=DOCX_TEXT_WHITE, size_pt=9.5)

    document.add_paragraph()

    # Secciones
    for sec_title, sec_index, is_further, blocks in parsed_sections:
        sec_cell = _docx_add_card(
            document,
            Inches,
            WD_TABLE_ALIGNMENT,
            WD_CELL_VERTICAL_ALIGNMENT,
            OxmlElement,
            qn,
            content_w_in,
            DOCX_BG_CARD_ALT,
            border_hex=None,
            accent_hex=DOCX_TEAL,
        )
        sec_cell.text = ""
        sec_p = sec_cell.paragraphs[0]
        _docx_style_paragraph(sec_p, Pt, Inches, after=0, line=1.1)
        prefix = f"{sec_index:02d}  " if (sec_index and not is_further) else ""
        _docx_add_bold_runs(
            sec_p,
            f"{prefix}{sec_title}",
            Pt,
            RGBColor,
            DOCX_TEAL,
            bold_color_hex=DOCX_TEAL,
            size_pt=12.5,
            force_bold=True,
        )
        document.add_paragraph()

        for block in blocks:
            kind = block[0]
            if kind == "body":
                p = document.add_paragraph()
                _docx_style_paragraph(p, Pt, Inches, after=6, line=1.25)
                _docx_add_bold_runs(p, block[1], Pt, RGBColor, DOCX_TEXT_GRAY, bold_color_hex=DOCX_TEXT_WHITE, size_pt=9.5)
                continue

            if kind == "table":
                rows_data = block[1]
                if rows_data:
                    table = document.add_table(rows=len(rows_data), cols=len(rows_data[0]))
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    for r_idx, row_cells in enumerate(rows_data):
                        row = table.rows[r_idx]
                        for c_idx, cell_text in enumerate(row_cells):
                            cell = row.cells[c_idx]
                            _docx_set_cell_background(cell, DOCX_BG_CARD_ALT if r_idx == 0 else DOCX_BG_CARD, OxmlElement, qn)
                            p = cell.paragraphs[0]
                            _docx_add_bold_runs(p, cell_text, Pt, RGBColor, DOCX_TEXT_WHITE, bold_color_hex=DOCX_TEXT_WHITE, size_pt=8, force_bold=(r_idx==0))
                document.add_paragraph()
                continue

            if kind == "image_ref":
                ref_spec = block[1]
                ref_name = ref_spec.get("name", "")
                image_path = ref_cache.get((p_num, ref_name))
                if image_path:
                    plan = _build_image_plan(image_path, ref_spec, content_w_in, image_max_h_in)
                    if plan["parts"] > 1 and PILImage is None:
                        log_fn(
                            f"    [ref] AVISO P{p_num}: Pillow no disponible, se omite split para '{ref_name}' en DOCX"
                        )
                        plan["parts"] = 1
                        plan["ranges"] = [(0.0, 1.0)]
                    for idx, (crop_start, crop_end) in enumerate(plan["ranges"]):
                        p = document.add_paragraph()
                        _docx_style_paragraph(
                            p,
                            Pt,
                            Inches,
                            align=WD_ALIGN_PARAGRAPH.CENTER,
                            after=4 if idx < len(plan["ranges"]) - 1 else 7,
                            line=1.0,
                        )
                        p.paragraph_format.keep_together = True
                        part_h = max(0.1, plan["draw_h"] * (crop_end - crop_start))
                        if len(plan["ranges"]) > 1:
                            crop_stream = _crop_image_to_stream(image_path, crop_start, crop_end)
                            if crop_stream is not None:
                                pic = p.add_run().add_picture(
                                    crop_stream,
                                    width=Inches(plan["draw_w"]),
                                    height=Inches(part_h),
                                )
                                _docx_style_inline_picture(pic, OxmlElement)
                            else:
                                # Fallback seguro si Pillow no está disponible.
                                pic = p.add_run().add_picture(image_path, width=Inches(plan["draw_w"]))
                                _docx_style_inline_picture(pic, OxmlElement)
                        else:
                            pic = p.add_run().add_picture(
                                image_path,
                                width=Inches(plan["draw_w"]),
                                height=Inches(part_h),
                            )
                            _docx_style_inline_picture(pic, OxmlElement)
                else:
                    warn_cell = _docx_add_card(
                        document, Inches, WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT,
                        OxmlElement, qn, content_w_in, DOCX_BG_SURFACE, border_hex=DOCX_WARN
                    )
                    warn_cell.text = ""
                    wp = warn_cell.paragraphs[0]
                    _docx_style_paragraph(wp, Pt, Inches, after=0, line=1.2)
                    _docx_add_bold_runs(
                        wp,
                        f"Imagen no encontrada para referencia: ref: {ref_name}",
                        Pt, RGBColor, DOCX_WARN, bold_color_hex=DOCX_WARN, size_pt=8.7
                    )
                document.add_paragraph()
                continue

            if kind == "image_ref_invalid":
                warn_cell = _docx_add_card(
                    document, Inches, WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT,
                    OxmlElement, qn, content_w_in, DOCX_BG_SURFACE, border_hex=DOCX_WARN
                )
                warn_cell.text = ""
                wp = warn_cell.paragraphs[0]
                _docx_style_paragraph(wp, Pt, Inches, after=0, line=1.2)
                _docx_add_bold_runs(
                    wp,
                    "Referencia de imagen vacia o invalida. Usa: ref: nombre_imagen",
                    Pt, RGBColor, DOCX_WARN, bold_color_hex=DOCX_WARN, size_pt=8.7
                )
                document.add_paragraph()
                continue

            if kind in ("case_ref", "prereq", "steps", "question", "example", "further", "closing"):
                if kind == "case_ref":
                    bg_hex, border_hex, accent_hex = DOCX_BG_CARD, DOCX_BLUE_LIGHT, None
                elif kind == "prereq":
                    bg_hex, border_hex, accent_hex = DOCX_BG_SURFACE, DOCX_WARN, None
                elif kind == "steps":
                    bg_hex, border_hex, accent_hex = DOCX_BG_CARD, None, None
                elif kind == "question":
                    bg_hex, border_hex, accent_hex = DOCX_BG_SURFACE, DOCX_TEAL_DIM, None
                elif kind == "example":
                    bg_hex, border_hex, accent_hex = DOCX_BG_CARD_ALT, DOCX_SUCCESS, None
                elif kind == "further":
                    bg_hex, border_hex, accent_hex = DOCX_BG_CARD, DOCX_BLUE_LIGHT, None
                else:  # closing
                    bg_hex, border_hex, accent_hex = DOCX_BG_CARD_ALT, None, DOCX_SUCCESS

                # Las cajas de listas en DOCX se dibujan un poco mas angostas para evitar efecto "bloque cuadrado".
                box_width_in = content_w_in - 0.14 if kind in ("prereq", "steps", "question", "example", "further") else content_w_in
                box_cell = _docx_add_card(
                    document,
                    Inches,
                    WD_TABLE_ALIGNMENT,
                    WD_CELL_VERTICAL_ALIGNMENT,
                    OxmlElement,
                    qn,
                    box_width_in,
                    bg_hex,
                    border_hex=border_hex,
                    accent_hex=accent_hex,
                )
                box_cell.text = ""
                base_p = box_cell.paragraphs[0]
                _docx_style_paragraph(base_p, Pt, Inches, after=0, line=1.2)

                if kind == "case_ref":
                    _docx_add_bold_runs(base_p, block[1], Pt, RGBColor, DOCX_TEXT_GRAY, bold_color_hex=DOCX_TEXT_WHITE, size_pt=9.5)
                elif kind == "prereq":
                    items = block[1]
                    base_p.text = ""
                    for idx, item in enumerate(items):
                        _docx_add_badge_row(
                            box_cell,
                            str(idx + 1),
                            item,
                            DOCX_TEXT_GRAY,
                            DOCX_TEXT_WHITE,
                            DOCX_TEAL,
                            DOCX_BG_PAGE,
                            DOCX_BG_SURFACE,
                            OxmlElement,
                            qn,
                            Inches,
                            Pt,
                            RGBColor,
                            WD_ALIGN_PARAGRAPH,
                            badge_w_in=0.22,
                            text_w_in=box_width_in - 0.24,
                            text_size_pt=9,
                        )
                elif kind == "steps":
                    steps = block[1]
                    base_p.text = ""
                    for sidx, step in enumerate(steps):
                        snum = step.get("snum", "")
                        marker = "•" if snum == "•" else str(snum)
                        text_cell = _docx_add_badge_row(
                            box_cell,
                            marker,
                            step["text"],
                            DOCX_TEXT_GRAY,
                            DOCX_TEXT_WHITE,
                            DOCX_TEAL,
                            DOCX_BG_PAGE,
                            DOCX_BG_CARD,
                            OxmlElement,
                            qn,
                            Inches,
                            Pt,
                            RGBColor,
                            WD_ALIGN_PARAGRAPH,
                            badge_w_in=0.22,
                            text_w_in=box_width_in - 0.24,
                            text_size_pt=9,
                        )
                        for sub in step.get("subs", []):
                            sub_text = sub.get("text", "") if isinstance(sub, dict) else str(sub)
                            sub_snum = sub.get("snum", "–") if isinstance(sub, dict) else "–"
                            sp = text_cell.add_paragraph()
                            _docx_style_paragraph(sp, Pt, Inches, after=1.2, line=1.14, left_indent=0.20, right_indent=0.12)
                            sub_prefix = "• " if sub_snum == "•" else (f"{sub_snum}. " if isinstance(sub_snum, int) else "– ")
                            _docx_add_bold_runs(sp, f"{sub_prefix}{sub_text}", Pt, RGBColor, DOCX_TEXT_DIM, bold_color_hex=DOCX_TEXT_GRAY, size_pt=8.5)
                elif kind == "question":
                    title_txt = block[1].strip().rstrip(':')
                    items = block[2]
                    start_idx = 0
                    if title_txt:
                        _docx_add_bold_runs(base_p, title_txt, Pt, RGBColor, DOCX_TEAL, bold_color_hex=DOCX_TEAL, size_pt=8.5, force_bold=True)
                    else:
                        base_p.text = ""
                        start_idx = 1
                    for qidx, item in enumerate(items):
                        p = box_cell.add_paragraph() if (title_txt or qidx > 0 or start_idx == 1) else base_p
                        _docx_style_paragraph(p, Pt, Inches, after=2, line=1.2, right_indent=0.12)
                        _docx_add_bold_runs(p, f"✓ {item}", Pt, RGBColor, DOCX_TEXT_GRAY, bold_color_hex=DOCX_TEXT_WHITE, size_pt=9)
                elif kind == "example":
                    title_txt = block[1].strip().rstrip(':')
                    items = block[2]
                    numbered = block[3]
                    if title_txt:
                        _docx_add_bold_runs(base_p, title_txt, Pt, RGBColor, DOCX_SUCCESS, bold_color_hex=DOCX_SUCCESS, size_pt=8.5, force_bold=True)
                    else:
                        base_p.text = ""
                    for eidx, item in enumerate(items):
                        p = box_cell.add_paragraph()
                        _docx_style_paragraph(p, Pt, Inches, after=2, line=1.2, right_indent=0.12)
                        prefix = f"{eidx+1}. " if numbered else "• "
                        _docx_add_bold_runs(p, f"{prefix}{item}", Pt, RGBColor, DOCX_TEXT_GRAY, bold_color_hex=DOCX_TEXT_WHITE, size_pt=9)
                elif kind == "further":
                    for fidx, item in enumerate(block[1]):
                        p = base_p if fidx == 0 else box_cell.add_paragraph()
                        _docx_style_paragraph(p, Pt, Inches, after=2, line=1.2, right_indent=0.12)
                        _docx_add_bold_runs(p, f"→ {item}", Pt, RGBColor, DOCX_BLUE_LIGHT, bold_color_hex=DOCX_TEXT_WHITE, size_pt=9)
                else:  # closing
                    _docx_add_bold_runs(base_p, block[1], Pt, RGBColor, DOCX_TEXT_GRAY, bold_color_hex=DOCX_TEXT_WHITE, size_pt=9)

                document.add_paragraph()
                continue

    base_no_ext = os.path.splitext(os.path.basename(md_path))[0]
    out_path = os.path.join(out_dir, f"{base_no_ext}.docx")
    document.save(out_path)
    return out_path

# === Añadido para UI runner: pipeline sin traducción ===
def run_pipeline(md_dir='', md_file='', logo_path='', out_dir='', p_from=1, p_to=999, output_format='pdf', log_fn=print):
    global LOGO_PATH
    if md_file:
        md_dir = os.path.dirname(os.path.abspath(md_file))
        out_dir = out_dir or md_dir
        md_files = []
        m = re.match(r'^P(\d+)', os.path.basename(md_file), re.IGNORECASE)
        if m:
            md_files = [(int(m.group(1)), md_file)]
    else:
        md_dir = (md_dir or '').rstrip('/')
        out_dir = out_dir or md_dir
        md_files = []
        for fname in sorted(os.listdir(md_dir)):
            m = re.match(r'^P(\d+)_.*\.md$', fname, re.IGNORECASE)
            if m:
                n = int(m.group(1))
                if p_from <= n <= p_to:
                    md_files.append((n, os.path.join(md_dir, fname)))

    if not md_dir or not os.path.isdir(md_dir):
        raise FileNotFoundError(f"Carpeta inválida: {md_dir}")
    os.makedirs(out_dir, exist_ok=True)

    output_format = (output_format or 'pdf').strip().lower()
    if output_format not in OUTPUT_FORMATS:
        raise RuntimeError(f"Formato inválido: {output_format}. Usa: {', '.join(OUTPUT_FORMATS)}")

    LOGO_PATH = ''
    for candidate in [logo_path,
                      os.path.join(md_dir, 'goal-logo-white.png'),
                      os.path.join(os.path.dirname(md_dir), 'goal-logo-white.png')]:
        if candidate and os.path.exists(candidate):
            LOGO_PATH = candidate
            break
    if not LOGO_PATH:
        log_fn("AVISO: goal-logo-white.png no encontrado — salida sin logotipo en encabezado.")

    if not md_files:
        raise RuntimeError("No se encontraron archivos Pxx_*.md en la carpeta.")

    log_fn(f"{len(md_files)} archivo(s) encontrados. Generando {output_format.upper()}...")
    ok, fail = 0, []
    for p_num, md_path in md_files:
        fname = os.path.basename(md_path)
        try:
            if output_format == 'pdf':
                out_path = build_pdf(md_path, out_dir, log_fn=log_fn)
            else:
                out_path = build_docx(md_path, out_dir, log_fn=log_fn)
            log_fn(f"  ✓ {os.path.basename(out_path)}")
            ok += 1
        except Exception as e:
            log_fn(f"  ✗ {fname}: {e}")
            fail.append(fname)

    if fail:
        log_fn(f"Fallaron: {', '.join(fail)}")
    log_fn(f"Concluido: {ok}/{len(md_files)}. Formato: {output_format.upper()}. Salida: {out_dir}")
    return {"ok": ok, "total": len(md_files), "out_dir": out_dir, "format": output_format}


# ── Entry point ──────────────────────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(description='goalbus Quick Start PDF generator (PT)')
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument('--md-dir', help='Carpeta con archivos Pxx_*.md en portugués')
    group.add_argument('--md', help='Archivo markdown individual')
    parser.add_argument('--logo', default='', help='Ruta a goal-logo-white.png')
    parser.add_argument('--out', default='', help='Carpeta de salida')
    parser.add_argument('--from', dest='p_from', type=int, default=1)
    parser.add_argument('--to', dest='p_to', type=int, default=99)
    parser.add_argument('--format', dest='output_format', default='pdf', choices=OUTPUT_FORMATS,
                        help='Formato de salida: pdf o docx')
    args = parser.parse_args()

    summary = run_pipeline(md_dir=args.md_dir, md_file=args.md, logo_path=args.logo,
                           out_dir=args.out, p_from=args.p_from, p_to=args.p_to,
                           output_format=args.output_format, log_fn=print)
    print(f"Resumen: {summary['ok']}/{summary['total']} ok. Formato: {summary['format'].upper()}. Salida: {summary['out_dir']}")


if __name__ == '__main__':
    main()
